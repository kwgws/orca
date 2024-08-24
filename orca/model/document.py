"""Model to track documents, anything to do with ORCA's corpus.
"""

import json
import logging
from pathlib import Path

from dateutil.parser import parse as parse_datetime
from docx import Document as Docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from sqlalchemy import Column, DateTime, ForeignKey, Integer, String
from sqlalchemy.orm import relationship
from unidecode import unidecode

from orca import config
from orca._helpers import utc_old
from orca.model.base import (
    Base,
    CommonMixin,
    documents_corpuses,
    documents_searches,
    with_session,
)

log = logging.getLogger(__name__)


class Image(Base, CommonMixin):
    """Each document is represented first by an immutable image file. These
    won't change as the document itself is revised.

    We need to be able to make exceptions, though, for documents which might
    come to us _without_ an image file; that might be imported eg as text only.
    Nevertheless, for organizational purposes, this is where we'll store as the
    base object.
    """

    album_index = Column(Integer, nullable=False)
    stem = Column(String(255), nullable=False)
    album = Column(String(255), nullable=False)
    title = Column(String(255), default="")
    media_archive = Column(String(255), default="")
    media_collection = Column(String(255), default="")
    media_box = Column(String(255), default="")
    media_folder = Column(String(255), default="")
    media_type = Column(String(255), default="")
    media_created_at = Column(DateTime, default=utc_old)
    scan_created_at = Column(DateTime, default=utc_old)
    image_path = Column(String(255), default="")
    image_url = Column(String(255), default="")
    thumb_url = Column(String(255), default="")
    documents = relationship(
        "Document", back_populates="image", cascade="all, delete-orphan"
    )

    @classmethod
    @with_session
    def create_from_file(cls, path, session=None, batch_only=False):
        return Document.create_from_file(
            path, None, session=session, batch_only=batch_only
        )

    def add_document_from_file(self, path, session=None, batch_only=False):
        return Document.create_from_file(
            path, self, session=session, batch_only=batch_only
        )

    def as_dict(self):
        rows = super().as_dict()
        rows.pop("image_path")
        if docs := [doc.as_dict() for doc in self.documents]:
            keys = set(docs[0].keys())
            keys_to_remove = {
                "stem",
                "title",
                "album",
                "album_index",
                "media_archive",
                "media_collection",
                "media_box",
                "media_folder",
                "media_type",
                "media_created_at",
                "image_url",
                "thumb_url",
            }
            for doc in docs:
                for key in keys.intersection(keys_to_remove):
                    doc.pop(key)
            rows["documents"] = docs
        return rows


class Document(Base, CommonMixin):
    """Documents here represent rather a specific revision of a given document.
    This is useful if we want to change something later: if we want to revise
    the text, or run the image through a different model, etc.

    A specific collection of documents will be stored in a Corpus. This should
    eventually let us do versioning and diff with our queries.
    """

    batch = Column(String(255), nullable=False)
    json_path = Column(String(255), default="")
    json_url = Column(String(255), default="")
    text_path = Column(String(255), default="")
    text_url = Column(String(255), default="")
    image_uid = Column(String(22), ForeignKey("images.uid"), nullable=False)
    image = relationship("Image", back_populates="documents")
    corpuses = relationship(
        "Corpus", back_populates="documents", secondary=documents_corpuses
    )
    searches = relationship(
        "Search", back_populates="documents", secondary=documents_searches
    )

    @property
    def stem(self):
        return self.image.stem

    @property
    def title(self):
        return self.image.title

    @property
    def album(self):
        return self.image.album

    @property
    def album_index(self):
        return self.image.album_index

    @property
    def media_archive(self):
        return self.image.media_archive

    @property
    def media_collection(self):
        return self.image.media_collection

    @property
    def media_box(self):
        return self.image.media_box

    @property
    def media_folder(self):
        return self.image.media_folder

    @property
    def media_type(self):
        return self.image.media_type

    @property
    def media_created_at(self):
        return self.image.media_created_at

    @property
    def scan_created_at(self):
        return self.image.scan_created_at

    @property
    def image_path(self):
        return self.image.image_path

    @property
    def image_url(self):
        return self.image.image_url

    @property
    def thumb_url(self):
        return self.image.thumb_url

    @property
    def json(self):
        try:
            with (config.data_path / self.json_path).open() as f:
                return json.load(f)
        except IOError:
            log.exception(f"Error loading file: {self.json_path}")
            raise

    @property
    def content(self):
        try:
            with (config.data_path / self.text_path).open() as f:
                return unidecode(f.read().strip())
        except IOError:
            log.exception(f"Error loading file: {self.text_path}")
            raise

    @classmethod
    @with_session
    def create_from_file(cls, path, image: Image, session=None, batch_only=False):
        """Commit a new Image/Document to the database. Use `batch_only` to
        indicate we only want to add to our current `session` rather than
        committing outright.

        This is based on a very specific kind of filename, which typically
        looks like this: `000001_2022-09-27_13-12-42_image_5992.json`. The data
        is separated by underscores. The first part is the index. The second
        two are the date and time. Any remaining parts are the title.
        """

        if not (path := Path(path)).is_file():
            raise FileNotFoundError(f"File not found: {path}")

        album = path.parent.name
        stem = path.stem
        split = stem.split("_")
        if len(split) < 3:
            raise ValueError(f"Invalid filename format: {stem}")

        # Parse the datetime from the filename using dateutil
        try:
            timestamp = parse_datetime(f"{split[1]} {split[2].replace('-', ':')}")
        except Exception as e:
            raise ValueError(f"Error parsing timestamp from filename: {stem}") from e

        # These paths need to be relative so we can make them portable
        json_path = Path(config.batch_name) / "json" / album / f"{stem}.json"
        text_path = Path(config.batch_name) / "text" / album / f"{stem}.txt"
        image_path = Path("img") / album / f"{stem}.webp"

        if not image or not isinstance(image, Image):
            image = Image(
                album_index=int(split[0]),
                stem=stem,
                title="_".join(split[3:]),
                album=album,
                image_path=f"{image_path}",
                image_url=f"{config.cdn.url}/{image_path}",
                thumb_url=f"{config.cdn.url}/thumbs/{album}/{stem}.webp",
                scan_created_at=timestamp,
            )

        document = Document(
            batch=config.batch_name,
            json_path=f"{json_path}",
            json_url=f"{config.cdn.url}/{json_path}",
            text_path=f"{text_path}",
            text_url=f"{config.cdn.url}/{text_path}",
        )

        document.image = image
        image.documents.append(document)
        session.add_all([image, document])
        if not batch_only:
            session.commit()
        return image

    def to_markdown_file(self, path: Path):
        """Appends content to a markdown file with metadata."""
        try:
            path.parent.mkdir(parents=True, exist_ok=True)
            with path.open("a") as f:
                f.write(
                    (
                        "---\n"
                        f"date: {self.created_at.strftime('%B %d, %Y at %-I:%M %p')}\n"
                        f"album: {self.title} - {self.album_index} of {self.album}\n"
                        f"image: {self.image_url}\n"
                        "---\n"
                        "\n"
                        f"{self.content}\n"
                        "\n"
                        "\n"
                        "\n"
                    )
                )

        except IOError:
            log.exception(f"Error writing to file: {path}")
            raise

    def to_docx_file(self, path: Path):
        """Generate or update a .DOCX file with text content and metadata."""
        try:
            # Open existing .DOCX file or create a new one
            path.parent.mkdir(parents=True, exist_ok=True)
            x = Docx(path.as_posix()) if path.exists() else Docx()

            # Add heading with metadata
            x.add_heading(self.created_at.strftime("%B %d, %Y at %-I:%M %p"), level=1)
            p = x.add_paragraph()
            run = p.add_run()
            run.text = f"{self.title} - {self.album_index} of {self.album}\n"
            run.font.bold = True

            # Add a link to the image file. To do this we need to manipulate
            # the underlying XML directly.
            run = OxmlElement("w:r")

            link = OxmlElement("w:hyperlink")  # Create link to image URL
            link.set(
                qn("r:id"),
                x.part.relate_to(
                    self.image_url,
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",  # noqa: E501
                    is_external=True,
                ),
            )

            # Create and style the link
            rPr = OxmlElement("w:rPr")
            color = OxmlElement("w:color")  # Blue
            color.set(qn("w:val"), "0000FF")
            rPr.append(color)
            underline = OxmlElement("w:u")
            underline.set(qn("w:val"), "single")  # Underlined
            rPr.append(underline)
            bold = OxmlElement("w:b")
            rPr.append(bold)  # Bold
            run.append(rPr)

            text_tag = OxmlElement("w:t")  # Set text value to URL
            text_tag.text = self.image_url
            run.append(text_tag)

            link.append(run)  # Add to paragraph
            p._p.append(link)  # Done!

            # Add OCRed text content followed by page break
            x.add_paragraph("-----")
            x.add_paragraph(self.content)
            x.add_page_break()

            # Save and move on. Phew!
            x.save(path)

        except Exception:
            log.exception(f"Unexpected error while creating .DOCX file: {path}")
            raise

    def as_dict(self):
        rows = super().as_dict()

        # Get rid of redundant image_uid and any local paths
        for key in {"image_uid", "json_path", "text_path"}:
            rows.pop(key)

        # Even though the Image class is storing documents for logical reasons,
        # most of our operations will happen on Documents directly. We should
        # expect this to be the most common interface for Image metadata.
        rows.update(
            {
                "stem": self.stem,
                "title": self.title,
                "album": self.album,
                "album_index": self.album_index,
                "media_archive": self.media_archive,
                "media_collection": self.media_collection,
                "media_box": self.media_box,
                "media_folder": self.media_folder,
                "media_type": self.media_type,
                "media_created_at": self.media_created_at,
                "image_url": self.image_url,
                "thumb_url": self.thumb_url,
            }
        )
        return rows
