"""
Module for generating and managing 'Megadoc' files from OCR-ed documents.

This module provides utilities for creating, saving, and uploading 'Megadoc'
files in various formats such as DOCX and Markdown. It integrates with
asynchronous file operations and S3-compatible storage for handling large
document collections efficiently.
"""

import asyncio
import logging
import mimetypes
from pathlib import Path
from random import random

import aioboto3
import aiofiles
from botocore.exceptions import BotoCoreError
from docx import Document as Docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from sqlalchemy.ext.asyncio import AsyncSession

from orca import config
from orca.model import Document, Megadoc, Search, with_async_session

log = logging.getLogger(__name__)


def _to_markdown_file(
    doc: Document,
    megadoc: Megadoc,
    is_last_page: bool = False,
    data_path: Path = config.data_path,
) -> None:
    """Creates and appends content to a markdown file.

    This function generates a markdown file using metadata from the provided
    `Document` and appends it to a file at the specified path. The content
    includes formatted text and metadata like the scan date, title, and album
    information.

    Parameters:
        doc (Document): The document object containing scan metadata.
        megadoc (Megadoc): The megadoc instance where the file path is stored.
        is_last_page (bool, optional): If `True`, omit the page break at the
            end of the text content. Defaults to `False`.
        data_path (Path, optional): Base data path where metadata files are
            stored. This is usually provided by `config.data_path` but can be
            overridden here for edge cases or testing.
    """
    path = data_path / megadoc.path
    path.parent.mkdir(parents=True, exist_ok=True)

    content = (
        "---\n"
        f"date: {doc.scan.scanned_at.strftime('%B %d, %Y at %-I:%M %p')}\n"
        f"album: {doc.scan.title} - {doc.scan.album_index} of {doc.scan.album}\n"
        f"image: {doc.scan.url}\n"
        "---\n"
        "\n"
        f"{doc.get_text(data_path=data_path)}\n"
        f"{'\n\n\n' if not is_last_page else ''}"
    )
    with path.open("a") as f:
        f.write(content)


def _to_docx_file(
    doc: Document,
    megadoc: Megadoc,
    is_last_page: bool = False,
    data_path: Path = config.data_path,
) -> None:
    """Creates a DOCX file from the given `Document` object, appending content
    with detailed metadata, OCR-ed text, and a hyperlink to the source image.

    This function manages the creation and styling of the DOCX document. It
    adds metadata as headings, inserts a hyperlink to an image file using low-
    level XML manipulation (due to limitations in the `python-docx` library),
    and includes OCR-ed text content. The document is saved to a specified
    path, which can be customized for testing or edge cases.

    Args:
        doc (Document): A `Document` instance containing metadata about the
            scanned document, including title, date, and URL to the image.
        megadoc (Megadoc): A `Megadoc` instance that holds the destination path
            where the DOCX file will be saved.
        is_last_page (bool, optional): If `True`, omit the page break at the
            end of the text content. Defaults to `False`.
        data_path (Path, optional): The base directory path where the DOCX
            files are stored. Defaults to `config.data_path` but can be
            overridden for testing or edge cases.
    """
    path = data_path / megadoc.path
    path.parent.mkdir(parents=True, exist_ok=True)
    x = Docx(str(path)) if path.exists() else Docx()

    x.add_heading(doc.scan.scanned_at.strftime("%B %d, %Y at %-I:%M %p"), level=1)
    p = x.add_paragraph()
    run = p.add_run()
    run.text = f"{doc.scan.title} - {doc.scan.album_index} of {doc.scan.album}\n"
    run.font.bold = True

    run = OxmlElement("w:r")
    link = OxmlElement("w:hyperlink")  # create link to image URL
    link.set(
        qn("r:id"),
        x.part.relate_to(
            doc.scan.url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",  # noqa: E501
            is_external=True,
        ),
    )

    rPr = OxmlElement("w:rPr")  # we need to manually style the link
    color = OxmlElement("w:color")  # blue
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)
    underline = OxmlElement("w:u")  # underlined
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    bold = OxmlElement("w:b")  # bold
    rPr.append(bold)
    run.append(rPr)

    text_tag = OxmlElement("w:t")
    text_tag.text = doc.scan.url  # type: ignore (implicit conversion)
    run.append(text_tag)

    link.append(run)  # add to paragraph
    p._p.append(link)

    x.add_paragraph("-----")
    x.add_paragraph(doc.get_text(data_path=data_path))
    if not is_last_page:
        x.add_page_break()

    x.save(str(path))


@with_async_session
async def create_megadoc(
    filetype: str,
    search: Search,
    *,
    data_path: Path = config.data_path,
    session: AsyncSession,
) -> Megadoc | None:
    """Asynchronously creates a megadoc from search results.

    This function generates a megadoc based on the specified `filetype` (e.g.,
    ".docx", ".txt") from the search results, including OCR-ed text and
    metadata. It tracks the progress and updates the megadoc's status
    accordingly.

    Parameters:
        filetype (str): The type of file to create (e.g., '.docx', '.md').
        search (Search): The search object containing the documents to compile.
        data_path (Path, optional): Base data path where metadata files are
            stored. This is usually provided by `config.data_path` but can be
            overridden here for edge cases or testing.
        session (AsyncSession, optional): An active asynchronous database
            session. If not provided, the method will create and manage its own
            session.

    Returns:
        A new megadoc, the old megadoc if this search already has one with this
            filetype, or `None` if the search yielded no results.

    Raises:
        NotImplementedError: If an unsupported file type is requested.
    """
    filetype = filetype.lower()  # quick sanity check
    if search.document_count < 1:
        log.warning(
            "🚧 Skipping Search '%s' <%s>, no results", search.search_str, search.guid
        )
        return None

    megadocs: list[Megadoc] = await search.awaitable_attrs.megadocs
    if old_megadoc := next((md for md in megadocs if md.filetype == filetype), None):
        log.warning(
            "🚧 Skipping Search '%s' <%s>, already has Megadoc <%s> of type '%s'",
            search.search_str,
            search.guid,
            old_megadoc.guid,
            filetype,
        )
        return old_megadoc

    log.info(
        "✨ Creating megadoc of type '%s' for Search '%s' <%s>",
        filetype,
        search.search_str,
        search.guid,
    )
    megadoc: Megadoc = await search.add_megadoc(filetype, session=session)

    documents: list[Document] = await search.awaitable_attrs.documents
    for i, doc in enumerate(sorted(documents, key=lambda doc: doc.created_at)):
        is_last_page = not i + 1 < search.document_count
        if filetype == ".docx":
            await asyncio.to_thread(
                _to_docx_file, doc, megadoc, is_last_page, data_path
            )
        elif filetype in {".md", ".txt"}:
            await asyncio.to_thread(
                _to_markdown_file, doc, megadoc, is_last_page, data_path
            )
        else:
            raise NotImplementedError(f"Cannot create megadoc of type {filetype}")

        await megadoc.update(
            {
                "progress": float(i + 1) / float(search.document_count),
                "status": "STARTED",
            },
            session=session,
        )

    # Set status to "SENDING" to indicate we're ready for upload
    log.info(
        "🌸 Done creating Megadoc <%s> of type '%s' for Search '%s' <%s>",
        megadoc.guid,
        filetype,
        search.search_str,
        search.guid,
    )
    await megadoc.update({"progress": 100.0, "status": "SENDING"}, session=session)
    return megadoc


@with_async_session
async def upload_megadoc(
    megadoc: Megadoc, *, data_path: Path = config.data_path, session: AsyncSession
) -> None:
    """Asynchronously uploads a megadoc file to an S3-compatible storage.

    This function handles the process of uploading a megadoc to S3, including
    retries on failure, setting the correct MIME type, and ensuring the file is
    publicly accessible.

    Parameters:
        megadoc (Megadoc): The megadoc instance containing the path and
            metadata.
        data_path (Path, optional): Base data path where metadata files are
            stored. This is usually provided by `config.data_path` but can be
            overridden here for edge cases or testing.
        session (AsyncSession, optional): An active asynchronous database
            session. If not provided, the method will create and manage its own
            session.

    Raises:
        FileNotFoundError: Megadoc file is not found in the specified path.
        RuntimeError: Upload failed after exhausting retry attempts.
    """
    path = data_path / megadoc.path
    if not path.is_file():
        raise FileNotFoundError(f"File not found: {path}")
    log.info("📡 Uploading Megadoc <%s> at %s to %s", megadoc.guid, path, megadoc.url)

    guess = await asyncio.to_thread(mimetypes.guess_type, str(path))
    content_type = guess[0] or "application/octet-stream"

    s3_session = aioboto3.Session()
    async with s3_session.client(  # type: ignore (internal issue w/ aioboto3)
        service_name="s3",
        region_name=config.s3.region,
        endpoint_url=config.s3.endpoint,
        aws_access_key_id=config.s3.access_key,
        aws_secret_access_key=config.s3.secret_key,
    ) as s3_client:
        for attempt in range(1, config.db.retries + 2):
            try:
                async with aiofiles.open(path, "rb") as file_bytes:
                    await s3_client.upload_fileobj(
                        file_bytes,
                        config.s3.space,
                        megadoc.path,
                        ExtraArgs={
                            "ACL": "public-read",
                            "ContentType": content_type,
                            "ContentDisposition": "attachment",
                        },
                    )

            except (OSError, BotoCoreError) as e:
                if attempt <= config.db.retries:
                    sleep_time = attempt**2 + random()  # jitter
                    log.warning(
                        "🚧 Error uploading Megadoc <%s>, "
                        "retrying in %.2f seconds (attempt %d of %d)",
                        megadoc.guid,
                        sleep_time,
                        attempt,
                        config.db.retries,
                    )
                    await asyncio.sleep(sleep_time)

                else:
                    raise RuntimeError(
                        f"Failed uploading Megadoc <{megadoc.guid}> "
                        f"after {attempt} attempts"
                    ) from e

    log.info("🌸 Done uploading Megadoc <%s> to %s", megadoc.guid, megadoc.url)
    await megadoc.set_status("SUCCESS", session=session)
