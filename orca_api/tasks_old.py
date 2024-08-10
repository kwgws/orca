import mimetypes
from pathlib import Path

import boto3
import click
from botocore.exceptions import BotoCoreError
from celery import Celery, Task, chain, chord, group
from natsort import natsorted
from sqlalchemy.exc import SQLAlchemyError
from unidecode import unidecode
from whoosh import index
from whoosh.fields import ID, TEXT, Schema
from whoosh.qparser import FuzzyTermPlugin, QueryParser
from whoosh.writing import AsyncWriter
from docx import Document as Docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from orca_api import config
from orca_api.db import Document, Megadoc, Search, redis_client, with_session

log = config.get_logger(__name__)


class BaseTask(Task):
    autoretry_for = (SQLAlchemyError, BotoCoreError)

    def on_failure(self, exc, task_id, args, kwargs, einfo):
        log.error(f"Task {self.name} failed: {exc} - {einfo}")


celery = Celery(__name__, config_source=config.CELERY)


@celery.task(bind=True, base=BaseTask)
@with_session
def load_documents(self, path, session=None):
    # Load list of files; sort, count
    doc_path = Path(path).absolute()
    json_files = natsorted(doc_path.glob("*.json"))
    doc_total = len(json_files)
    log.info(f"Importing {doc_total} docuents from {doc_path}")

    # Get last processed index from Redis.
    redis_key = f"load_documents:{path}"
    doc_count = int(redis_client.hget(redis_key, "doc_count") or -1)

    # Start processing from next file.
    for i in range(doc_count + 1, doc_total):
        json_file = json_files[i]
        if (i + 1) % 1000 == 0 or i + 1 == doc_total:
            log.info(f"Importing {doc_path} ({i + 1}/{doc_total})")
            Document.create(json_file, session=session)
            redis_client.hset(redis_key, "doc_count", i)
        else:
            Document.create(json_file, session=session, no_commit=True)

    redis_client.delete(redis_key)
    return


@celery.task(bind=True, base=BaseTask)
@with_session
def index_documents(self, _, session=None):
    index_path = config.BATCH_PATH / "index"
    index_path.mkdir(parents=True, exist_ok=True)

    schema = Schema(
        id=ID(stored=True, unique=True),
        content=TEXT(stored=True),
    )

    ix = index.create_in(index_path, schema)
    writer = AsyncWriter(ix)

    log.info("Generating table hash")
    Document.update_hash(session=session)
    documents = Document.get_all(session=session)
    doc_total = Document.get_count(session=session)

    # Get last processed index from Redis.
    redis_key = f"build_index:{index_path}"
    doc_count = int(redis_client.hget(redis_key, "doc_count") or -1)
    log.info(f"Indexing {doc_total} docuents to {index_path}")

    for i in range(doc_count + 1, doc_total):
        doc = documents[i]
        txt_path = config.DATA_PATH / doc.txt_path

        try:
            with open(txt_path, encoding="utf-8") as f:
                content = unidecode(f.read().strip())
            writer.add_document(id=doc.id, content=content)

            if (i + 1) % 10000 == 0 or i + 1 == doc_total:
                log.info(f"Indexing documents to {index_path} ({i + 1}/{doc_total})")
                redis_client.hset(redis_key, "doc_count", i)

        except FileNotFoundError as e:
            log.warning(f"Error parsing {txt_path}: {e}")

    log.info("Finalizing index")
    writer.commit()
    redis_client.delete(redis_key)
    log.info("Done!")


@celery.task(bind=True, base=BaseTask)
@with_session
def process_search(self, search_id, session=None):
    log.info(f"Starting process_search() with id {search_id}")
    search = Search.get(search_id, session=session)
    if not search:
        log.warning(f"No search found with id {search_id}")
        return

    log.info(f"Processing `{search.search_str}` with id {search_id}")

    # Load whoosh index
    ix = index.open_dir(config.INDEX_PATH.as_posix())
    with ix.searcher() as searcher:
        parser = QueryParser("content", ix.schema)
        parser.add_plugin(FuzzyTermPlugin())

        # Parse query and perform search
        whoosh_query = parser.parse(search.search_str)
        for result in searcher.search(whoosh_query, limit=None):
            document = Document.get(result["id"], session=session)
            search.add_result(document, session=session)

    search.set_status("SUCCESS", session=session)
    return search_id


@celery.task(bind=True, base=BaseTask)
@with_session
def process_megadoc(self, search_id, filetype, session=None):
    log.debug(f"Starting process_megadoc with id {search_id} and filetype {filetype}")
    search = Search.get(search_id, session=session)
    if not search:
        log.warning(f"No search with id {search_id}")
        return

    megadoc = Megadoc.get_by_search(search, filetype, session=session)
    if not megadoc:
        log.debug(f"Creating new {filetype} for `{search.search_str}`")
        megadoc = Megadoc.create(search, filetype, session=session)
    elif megadoc.status == "SUCCESS":
        log.warning(f"Tried to re-run {filetype} for `{search.search_str}`")
        return megadoc.id

    documents = sorted(search.results, key=lambda d: d.created.isoformat())
    log.info(f"Processing {filetype} for `{search.search_str}`")
    for i in range(megadoc.get_progress(), len(documents)):
        megadoc.add_doc_to_megadoc(documents[i])
        if megadoc.status != "STARTED":
            megadoc.set_status("STARTED", session=session)

    megadoc.set_status("SENDING", session=session)
    return megadoc.id


@celery.task(bind=True, base=BaseTask)
@with_session
def upload_megadoc(self, megadoc_id, session=None):
    log.debug(f"Starting upload_megadoc with id {megadoc_id}")
    megadoc = Megadoc.get(megadoc_id, session=session)
    if not megadoc:
        log.warning(f"No megadoc with id {megadoc.id}")
        return
    search = Search.get(megadoc.search_id, session=session)

    log.info(f"Uploading {megadoc.filetype} for `{search.search_str}`")

    # Get document mime type
    local_path = (config.DATA_PATH / megadoc.path).as_posix()
    content_type, _ = mimetypes.guess_type(local_path)
    if content_type is None:
        content_type = "application/octet-stream"

    try:
        # Upload megadoc
        s3_session = boto3.session.Session()
        s3_client = s3_session.client(
            "s3",
            region_name=config.CDN_REGION,
            endpoint_url=config.CDN_ENDPOINT,
            aws_access_key_id=config.CDN_ACCESS_KEY,
            aws_secret_access_key=config.CDN_SECRET_KEY,
        )
        s3_client.upload_file(
            local_path,
            config.CDN_SPACE_NAME,
            megadoc.path,
            ExtraArgs={
                "ACL": "public-read",
                "ContentType": content_type,
                "ContentDisposition": "attachment",
            },
        )
    except BotoCoreError as e:
        log.error(f"Unexpected error uploading megadoc to S3: {e}")
        raise self.retry(exc=e)

    megadoc.set_status("SUCCESS", session=session)
    return megadoc.id


@click.command()
@click.argument("base_path", type=click.Path(exists=True, file_okay=False))
def start_load(base_path):
    base_path = Path(base_path)
    if not base_path.is_dir():
        click.echo(f"{base_path} is not a directory")
        return

    subdirs = [d for d in base_path.iterdir() if d.is_dir()]
    if not subdirs:
        click.echo(f"No subdirectories found in {base_path}")
        return

    import_tasks = [load_documents.s(p.as_posix()) for p in subdirs]
    result = chord(import_tasks)(index_documents.s())
    click.echo(f"Started import tasks for {len(subdirs)} subdirectories.")
    click.echo(f"{result}")


def start_search(search_id: str):
    log.info(f"Starting tasks for search with id {search_id}")

    workflow = chain(
        process_search.s(search_id),
        group(
            chain(process_megadoc.s(filetype), upload_megadoc.s())
            for filetype in config.MEGADOC_FILETYPES
        ),
    )

    result = workflow.apply_async()
    return result

def get_text(self):
        try:
            with self.txt_path.open() as f:
                content = f.read().strip()
            if not content or content == "":
                log.warning(f"No content: {self.txt_path}")
        except (FileNotFoundError, IsADirectoryError, PermissionError) as e:
            log.warning(f"File not found: {self.txt_path}: {e}")
        return content

    def to_markdown(self, path: Path):
        with path.open("a") as f:
            f.write(
                "\n".join(
                    "---",
                    f"date: {self.strftime('%B %d, %Y at %-I:%M %p')}",
                    f"album: {self.title} - {self.index} of {self.album}",
                    f"image: {self.img_url}",
                    "---",
                    "",
                    self.get_text(),
                    "",
                    "",
                    "",
                )
            )

    def to_docx(self, path: Path):
        x = Docx(path.as_posix()) if path.exists() else Docx()

        x.add_heading(self.created, level=1)

        p = x.add_paragraph()
        run = p.add_run()
        run.text = f"{self.title} - {self.index} of {self.album}\n"
        run.font.bold = True

        # To add a link we need to manipulate the underlying XML directly.
        run = OxmlElement("w:r")

        link = OxmlElement("w:hyperlink")  # Create link
        link.set(
            qn("r:id"),
            x.part.relate_to(
                self.img_url,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",  # noqa: E501
                is_external=True,
            ),
        )

        rPr = OxmlElement("w:rPr")  # Format
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0000FF")
        rPr.append(color)
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        rPr.append(underline)
        bold = OxmlElement("w:b")
        rPr.append(bold)
        run.append(rPr)

        text_tag = OxmlElement("w:t")  # Set text
        text_tag.text = self.img_url
        run.append(text_tag)

        link.append(run)  # Add to paragraph
        p._p.append(link)  # Done!

        x.add_paragraph("-----")
        x.add_paragraph(self.get_text())
        x.add_page_break()

        x.save(path)
